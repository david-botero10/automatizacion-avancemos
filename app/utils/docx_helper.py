"""
Utilidades para manipular documentos de Word (.docx).
Facilita operaciones comunes como reemplazo de texto y guardado de documentos.
"""

import os
import logging
from docx import Document

# Configurar logger para este módulo
logger = logging.getLogger(__name__)

def replace_text_in_doc(doc, reemplazos):
    """
    Reemplaza texto en un documento Word según una lista de reemplazos.
    
    Args:
        doc (Document): Documento de Word abierto
        reemplazos (list): Lista de diccionarios con claves 'original' y 'nuevo'
    
    Returns:
        bool: True si se realizó al menos un reemplazo, False en caso contrario
    """
    if not doc or not reemplazos:
        return False
        
    reemplazos_realizados = 0
    
    # Procesar párrafos
    for i, paragraph in enumerate(doc.paragraphs):
        for reemplazo in reemplazos:
            original = reemplazo.get('original')
            nuevo = reemplazo.get('nuevo')
            
            if original and nuevo and original in paragraph.text:
                # Reemplazar en el texto del párrafo
                paragraph.text = paragraph.text.replace(original, nuevo)
                reemplazos_realizados += 1
                logger.debug(f"Reemplazo realizado en párrafo {i}: {original} -> {nuevo}")
    
    # Procesar tablas
    for i, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for reemplazo in reemplazos:
                        original = reemplazo.get('original')
                        nuevo = reemplazo.get('nuevo')
                        
                        if original and nuevo and original in paragraph.text:
                            # Reemplazar en el texto del párrafo
                            paragraph.text = paragraph.text.replace(original, nuevo)
                            reemplazos_realizados += 1
                            logger.debug(f"Reemplazo realizado en tabla {i}: {original} -> {nuevo}")
    
    logger.info(f"Total de reemplazos realizados: {reemplazos_realizados}")
    return reemplazos_realizados > 0

def save_document(doc, ruta_destino):
    """
    Guarda un documento Word en la ruta especificada.
    Crea directorios intermedios si no existen.
    
    Args:
        doc (Document): Documento de Word a guardar
        ruta_destino (str): Ruta completa donde guardar el documento
    
    Returns:
        bool: True si se guardó correctamente, False en caso contrario
        
    Raises:
        Exception: Si ocurre algún error al guardar el documento
    """
    try:
        # Crear directorio si no existe
        directorio = os.path.dirname(ruta_destino)
        if directorio and not os.path.exists(directorio):
            os.makedirs(directorio, exist_ok=True)
            logger.info(f"Creado directorio: {directorio}")
        
        # Guardar documento
        doc.save(ruta_destino)
        logger.info(f"Documento guardado en: {ruta_destino}")
        return True
        
    except Exception as e:
        logger.error(f"Error al guardar documento en {ruta_destino}: {str(e)}")
        raise
        
def extract_text_from_doc(ruta_archivo):
    """
    Extrae todo el texto de un documento Word.
    
    Args:
        ruta_archivo (str): Ruta al archivo .docx
    
    Returns:
        str: Texto completo del documento
        
    Raises:
        FileNotFoundError: Si el archivo no existe
        Exception: Si ocurre algún error al abrir o procesar el documento
    """
    if not os.path.exists(ruta_archivo):
        logger.error(f"Archivo no encontrado: {ruta_archivo}")
        raise FileNotFoundError(f"Archivo no encontrado: {ruta_archivo}")
        
    try:
        doc = Document(ruta_archivo)
        texto_completo = "\n".join([p.text for p in doc.paragraphs])
        
        # También extraer texto de tablas
        for table in doc.tables:
            for row in table.rows:
                texto_fila = " | ".join([cell.text for cell in row.cells])
                texto_completo += f"\n{texto_fila}"
                
        logger.debug(f"Extraído texto de {ruta_archivo} ({len(texto_completo)} caracteres)")
        return texto_completo
        
    except Exception as e:
        logger.error(f"Error al extraer texto de {ruta_archivo}: {str(e)}")
        raise