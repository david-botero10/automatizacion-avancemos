"""
Módulo principal para el procesamiento de expedientes.
Contiene la clase ProcesadorExpedientes que maneja la lógica de negocio.
"""

import os
import re
import json
import logging
from docx import Document
import traceback
from datetime import datetime

# Importar utilidades propias
try:
    from .utils.docx_helper import replace_text_in_doc, save_document
    from .utils.logger import setup_logger
except ImportError:
    # En caso de ejecutarse directamente
    from utils.docx_helper import replace_text_in_doc, save_document
    from utils.logger import setup_logger

class ProcesadorExpedientes:
    """
    Clase principal para procesar expedientes de insolvencia.
    Maneja la extracción de información y generación de documentos.
    """
    
    def __init__(self, config):
        """
        Inicializa el procesador de expedientes.
        
        Args:
            config (dict): Configuración con rutas y parámetros
        """
        self.ruta_base = config.get('ruta_expedientes', '')
        self.ruta_formatos = config.get('ruta_formatos', '')
        self.ruta_salida = config.get('ruta_salida', self.ruta_base)
        
        # Configurar logger
        self.logger = setup_logger(
            nombre="procesador", 
            nivel=config.get('nivel_log', 'INFO'),
            ruta_log=config.get('ruta_log', os.path.join(os.path.dirname(os.path.dirname(__file__)), 'logs'))
        )
        
        # Cargar mapeo de operadores
        self.operadores_formatos = self._cargar_mapeo_operadores()
        
        self.logger.info(f"Procesador inicializado con {len(self.operadores_formatos)} operadores mapeados")
        
    def _cargar_mapeo_operadores(self):
        """
        Carga el mapeo de operadores desde el archivo JSON.
        Si no existe, crea un mapeo automático.
        
        Returns:
            dict: Diccionario con nombres de operadores como claves y rutas de formatos como valores.
        """
        # Determinar ruta del archivo JSON
        ruta_json = self._get_operadores_json_path()
        
        # Intentar cargar desde JSON
        if os.path.exists(ruta_json):
            try:
                with open(ruta_json, 'r', encoding='utf-8') as f:
                    operadores = json.load(f)
                    # Convertir rutas relativas a absolutas
                    for operador, ruta in operadores.items():
                        if not os.path.isabs(ruta):
                            operadores[operador] = os.path.join(self.ruta_formatos, ruta)
                    self.logger.info(f"Mapeo de operadores cargado desde {ruta_json}")
                    return operadores
            except Exception as e:
                self.logger.error(f"Error al cargar mapeo de operadores: {str(e)}")
        
        # Si no se pudo cargar, crear mapeo automático
        self.logger.warning("Creando mapeo automático de operadores")
        return self._mapear_operadores_formatos()
    
    def _get_operadores_json_path(self):
        """
        Determina la ruta del archivo JSON de operadores.
        
        Returns:
            str: Ruta al archivo JSON
        """
        # Posibles ubicaciones
        posibles_rutas = [
            os.path.join(os.path.dirname(__file__), 'config', 'operadores.json'),
            os.path.join(os.path.dirname(os.path.dirname(__file__)), 'config', 'operadores.json'),
            os.path.join(os.path.dirname(os.path.dirname(__file__)), 'data', 'operadores.json')
        ]
        
        # Verificar cada ruta
        for ruta in posibles_rutas:
            if os.path.exists(ruta):
                return ruta
        
        # Si no existe, usar la primera opción
        return posibles_rutas[0]
        
    def _mapear_operadores_formatos(self):
        """
        Crea un diccionario que mapea nombres de operadores a sus archivos de formato.
        
        Returns:
            dict: Diccionario con nombres de operadores como claves y rutas de formatos como valores.
        """
        operadores_formatos = {}
        
        # Verificar que la ruta de formatos exista
        if not os.path.exists(self.ruta_formatos):
            self.logger.error(f"La ruta de formatos no existe: {self.ruta_formatos}")
            return operadores_formatos
            
        for archivo in os.listdir(self.ruta_formatos):
            if archivo.endswith('.docx'):
                try:
                    doc = Document(os.path.join(self.ruta_formatos, archivo))
                    # Buscamos el nombre del operador en el contenido del documento
                    for paragraph in doc.paragraphs:
                        operador_match = re.search(r'([A-Z\s]{10,}GUERRERO|[A-Z\s]{10,})', paragraph.text)
                        if operador_match:
                            nombre_operador = operador_match.group(0).strip()
                            operadores_formatos[nombre_operador] = os.path.join(self.ruta_formatos, archivo)
                            self.logger.info(f"Mapeado operador '{nombre_operador}' a formato '{archivo}'")
                            break
                except Exception as e:
                    self.logger.error(f"Error al procesar formato {archivo}: {str(e)}")
        
        # Guardar el mapeo para futuras ejecuciones
        self._guardar_mapeo_operadores(operadores_formatos)
        
        return operadores_formatos
    
    def _guardar_mapeo_operadores(self, operadores_formatos):
        """
        Guarda el mapeo de operadores en un archivo JSON.
        
        Args:
            operadores_formatos (dict): Diccionario con mapeo de operadores a formatos.
        """
        ruta_json = self._get_operadores_json_path()
        try:
            # Crear directorio si no existe
            os.makedirs(os.path.dirname(ruta_json), exist_ok=True)
            
            # Convertir rutas absolutas a relativas para mayor portabilidad
            operadores_rel = {}
            for operador, ruta in operadores_formatos.items():
                if os.path.isabs(ruta) and ruta.startswith(self.ruta_formatos):
                    operadores_rel[operador] = os.path.basename(ruta)
                else:
                    operadores_rel[operador] = ruta
            
            with open(ruta_json, 'w', encoding='utf-8') as f:
                json.dump(operadores_rel, f, indent=4, ensure_ascii=False)
                
            self.logger.info(f"Guardado mapeo de operadores en {ruta_json}")
        except Exception as e:
            self.logger.error(f"Error al guardar mapeo de operadores: {str(e)}")
    
    def procesar_expedientes(self):
        """
        Procesa todos los expedientes en la ruta base, ignorando los que tienen '00' en el nombre.
        
        Returns:
            tuple: (expedientes_procesados, expedientes_ignorados, expedientes_error)
        """
        expedientes_procesados = 0
        expedientes_ignorados = 0
        expedientes_error = 0
        
        self.logger.info(f"Iniciando procesamiento de expedientes en {self.ruta_base}")
        
        # Verificar que la ruta base exista
        if not os.path.exists(self.ruta_base):
            self.logger.error(f"La ruta base no existe: {self.ruta_base}")
            return (0, 0, 0)
            
        for expediente in os.listdir(self.ruta_base):
            # Ignorar expedientes con '00' en el nombre
            if ' 00 ' in expediente:
                self.logger.info(f"Ignorando expediente con '00': {expediente}")
                expedientes_ignorados += 1
                continue
                
            ruta_expediente = os.path.join(self.ruta_base, expediente)
            if os.path.isdir(ruta_expediente):
                try:
                    if self.procesar_expediente(ruta_expediente):
                        expedientes_procesados += 1
                    else:
                        expedientes_error += 1
                except Exception as e:
                    self.logger.error(f"Error al procesar expediente {expediente}: {str(e)}")
                    expedientes_error += 1
        
        self.logger.info(f"Procesamiento finalizado. Procesados: {expedientes_procesados}, "
                         f"Ignorados: {expedientes_ignorados}, Errores: {expedientes_error}")
        
        return expedientes_procesados, expedientes_ignorados, expedientes_error
    
    def procesar_expediente(self, ruta_expediente):
        """
        Procesa un expediente individual.
        
        Args:
            ruta_expediente (str): Ruta del expediente a procesar.
            
        Returns:
            bool: True si el procesamiento fue exitoso, False en caso contrario.
        """
        nombre_expediente = os.path.basename(ruta_expediente)
        self.logger.info(f"Procesando expediente: {nombre_expediente}")
        
        # Verificar si existen las carpetas necesarias
        carpeta_principal = os.path.join(ruta_expediente, "01. CUADERNO PRINCIPAL")
        carpeta_notificaciones = os.path.join(ruta_expediente, "02. NOTIFICACIONES")
        
        if not os.path.exists(carpeta_principal):
            self.logger.warning(f"Carpeta '01. CUADERNO PRINCIPAL' no encontrada en {nombre_expediente}")
            return False
        
        if not os.path.exists(carpeta_notificaciones):
            self.logger.info(f"Carpeta '02. NOTIFICACIONES' no encontrada en {nombre_expediente}. Creándola.")
            try:
                os.makedirs(carpeta_notificaciones)
            except Exception as e:
                self.logger.error(f"Error al crear carpeta de notificaciones: {str(e)}")
                return False
        
        # Buscar archivo de aceptación de solicitud
        archivo_aceptacion = None
        for archivo in os.listdir(carpeta_principal):
            if archivo.startswith("Aceptación de solicitud"):
                archivo_aceptacion = os.path.join(carpeta_principal, archivo)
                break
        
        if not archivo_aceptacion:
            self.logger.warning(f"No se encontró archivo de aceptación en {nombre_expediente}")
            return False
        
        # Extraer información del archivo de aceptación
        try:
            info_deudor = self.extraer_informacion_aceptacion(archivo_aceptacion)
            if not info_deudor:
                return False
                
            # Generar notificación para acreedores
            return self.generar_notificacion_acreedores(info_deudor, carpeta_notificaciones)
            
        except Exception as e:
            self.logger.error(f"Error al procesar {archivo_aceptacion}: {str(e)}")
            self.logger.error(traceback.format_exc())
            return False
    
    def extraer_informacion_aceptacion(self, ruta_archivo):
        """
        Extrae información de un archivo de aceptación de solicitud.
        
        Args:
            ruta_archivo (str): Ruta al archivo de aceptación.
            
        Returns:
            dict: Diccionario con la información extraída del deudor.
        """
        self.logger.info(f"Extrayendo información de {os.path.basename(ruta_archivo)}")
        
        try:
            doc = Document(ruta_archivo)
            texto_completo = "\n".join([p.text for p in doc.paragraphs])
            
            # Extraer nombre del deudor
            nombre_deudor_match = re.search(r'Deudora?\s*\n*\s*([A-Z\s]+)\s*\n*\s*CC No', texto_completo)
            if not nombre_deudor_match:
                nombre_deudor_match = re.search(r'ELVIN CECILIA TORRES DURAN|[A-Z\s]{10,}', texto_completo)
            
            # Extraer cédula
            cedula_match = re.search(r'CC No\.\s*(\d+(?:\.\d+)*)', texto_completo)
            if not cedula_match:
                cedula_match = re.search(r'cédula de ciudadanía número\s*(\d+(?:\.\d+)*)', texto_completo)
            
            # Extraer radicado
            radicado_match = re.search(r'Radicado:\s*([0-9-]+)', texto_completo)
            
            # Extraer fecha de presentación y audiencia
            fecha_presentacion_match = re.search(r'presentó solicitud de negociación de sus deudas.*?el día (\d+ de [a-zA-Z]+ de \d+)', texto_completo)
            fecha_audiencia_match = re.search(r'audiencia de negociación de pasivos\s*.*?el día (\d+.*?\d+)', texto_completo)
            
            # Extraer operador de insolvencia (última página)
            operador_match = None
            for i in range(len(doc.paragraphs) - 1, 0, -1):
                match = re.search(r'DIANA PATRICIA MANGA GUERRERO|[A-Z\s]{10,}GUERRERO', doc.paragraphs[i].text)
                if match:
                    operador_match = match
                    break
                
            if not nombre_deudor_match or not cedula_match or not radicado_match or not operador_match:
                self.logger.warning(f"No se pudieron extraer todos los datos requeridos de {os.path.basename(ruta_archivo)}")
                missing_data = []
                if not nombre_deudor_match: missing_data.append("nombre_deudor")
                if not cedula_match: missing_data.append("cedula")
                if not radicado_match: missing_data.append("radicado")
                if not operador_match: missing_data.append("operador")
                self.logger.warning(f"Datos faltantes: {', '.join(missing_data)}")
                return None
                
            info = {
                'nombre_deudor': (nombre_deudor_match.group(1) if hasattr(nombre_deudor_match, 'group') and 
                                nombre_deudor_match.group(1) else nombre_deudor_match.group(0)).strip(),
                'cedula': cedula_match.group(1),
                'radicado': radicado_match.group(1),
                'operador': operador_match.group(0).strip(),
                'fecha_extraccion': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            }
            
            # Añadir fechas si están disponibles
            if fecha_presentacion_match:
                info['fecha_presentacion'] = fecha_presentacion_match.group(1)
            if fecha_audiencia_match:
                info['fecha_audiencia'] = fecha_audiencia_match.group(1)
                
            self.logger.info(f"Información extraída: {json.dumps(info, ensure_ascii=False)}")
            return info
                
        except Exception as e:
            self.logger.error(f"Error al extraer información de {os.path.basename(ruta_archivo)}: {str(e)}")
            self.logger.error(traceback.format_exc())
            return None
    
    def generar_notificacion_acreedores(self, info_deudor, carpeta_destino):
        """
        Genera una notificación para acreedores basada en el formato del operador.
        
        Args:
            info_deudor (dict): Información del deudor extraída del archivo de aceptación.
            carpeta_destino (str): Carpeta donde se guardará la notificación generada.
            
        Returns:
            bool: True si la generación fue exitosa, False en caso contrario.
        """
        operador = info_deudor['operador']
        
        # Buscar formato correspondiente al operador
        formato_path = None
        for nombre_operador, ruta_formato in self.operadores_formatos.items():
            if operador in nombre_operador or nombre_operador in operador:
                formato_path = ruta_formato
                break
        
        if not formato_path:
            self.logger.warning(f"No se encontró formato para el operador: {operador}")
            return False
        
        try:
            # Abrir formato
            doc = Document(formato_path)
            
            # Preparar reemplazos
            reemplazos = [
                {"original": "Señores", "nuevo": "Señor(a)"},
                {"original": "**Deudor:**", "nuevo": f"**Deudor:** {info_deudor['nombre_deudor']}"},
                {"original": "**C.C.**", "nuevo": f"**C.C.** {info_deudor['cedula']}"},
                {"original": "**Radicado:**", "nuevo": f"**Radicado:** {info_deudor['radicado']}"}
            ]
            
            # Añadir fechas si están disponibles
            if 'fecha_presentacion' in info_deudor:
                reemplazos.append({
                    "original": "el día **\\_\\_\\_\\_\\_\\_\\_\\_\\_\\_**", 
                    "nuevo": f"el día **{info_deudor['fecha_presentacion']}**"
                })
            
            if 'fecha_audiencia' in info_deudor:
                reemplazos.append({
                    "original": "el día **\\_\\_\\_\\_\\_\\_\\_\\_\\_\\_\\_\\_\\_\\_-**", 
                    "nuevo": f"el día **{info_deudor['fecha_audiencia']}**"
                })
            
            # Aplicar reemplazos
            replace_text_in_doc(doc, reemplazos)
            
            # Guardar documento modificado
            nombre_archivo = f"Notificación_{info_deudor['nombre_deudor']}.docx"
            ruta_salida = os.path.join(carpeta_destino, nombre_archivo)
            save_document(doc, ruta_salida)
            
            self.logger.info(f"Notificación generada exitosamente: {nombre_archivo}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error al generar notificación: {str(e)}")
            self.logger.error(traceback.format_exc())
            return False